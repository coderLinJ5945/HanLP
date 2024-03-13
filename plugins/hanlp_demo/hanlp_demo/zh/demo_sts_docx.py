# -*- coding:utf-8 -*-
# Author: hankcs
# Date: 2021-05-24 13:15
import hanlp

# sim = hanlp.load(hanlp.pretrained.sts.STS_ELECTRA_BASE_ZH)
# result = sim([
#     ['1、对于本项目的理解及重点、难点分析；', '1、对于本工程的理解及重点、难点理解'],
#     ['无线路由器怎么无线上网', '2、总体方案及思路'],
#     ['北京到上海的动车票', '上海到北京的动车票'],
# ])
# print(type(result))
# print(result)
###### pip start########

# pip install python-docx
###### pip end  ########
import re
from docx import Document


def extract_sentences_from_docx(doc_file):
    """
    读取docx文件，返回文件的句子字符串列表
    参数:
    docx_path (str): docx文件的路径
    返回:
    list[str]: 包含文件中所有句子的字符串列表
    """
    # 读取docx文件
    doc = Document(doc_file)

    pattern = re.compile(r'([^。！!？?;；\n]*)([。！!？?;；]|[\r\n]+)')

    # 将文档内容合并为一个字符串（按段落连接）
    # text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    # 按照中文句子进行拆分（这里简单地以句号、问号、感叹号结尾作为句子结束标志）
    # chinese_sentences = [match.group(0) for match in re.finditer(r'[^。！!？?;；]*[。！!？?;；]', text)]
    matches = pattern.finditer(text)
    chinese_sentences = [match.group(0) for match in matches]
    # 过滤掉空字符串
    # chinese_sentences = [sentence for sentence in chinese_sentences if sentence]
    # 过滤掉空字符串和小于3的，防止正则匹配到的符号传入到了模型中 todo 这里可以根据扩展根据需求进行过滤筛选，减少模型的比对次数
    chinese_sentences = [sentence for sentence in chinese_sentences if sentence and len(sentence.strip()) > 3]

    # 输出每一句，并统计句子数量
    # sentence_count = 0
    # for index, sentence in enumerate(chinese_sentences):
    #     # print(f"{sentence.strip()}")
    #     print(f"{index + 1}: {sentence.strip()}")
    #     sentence_count += 1
    #
    # print(f"总共识别到 {sentence_count} 句话")
    return chinese_sentences

# 调用函数
# extract_sentences_from_docx('C:\\Users\\linj\\Desktop\\tmp\\test\\test.docx')


def cross_match_sentences(sentences_list1, sentences_list2):
    """
    将第一个句子列表中的每个句子与第二个句子列表中的所有句子进行交叉匹配，
    生成包含所有可能组合的二维数组。

    参数:
    sentences_list1 (list[str]): 第一个文档的句子列表
    sentences_list2 (list[str]): 第二个文档的句子列表

    返回:
    list[list[str]]: 每个内部列表包含一对句子的二维数组
    """
    cross_matched_sentences = [[sent1, sent2] for sent1 in sentences_list1 for sent2 in sentences_list2]

    return cross_matched_sentences


if __name__ == '__main__':
    docx1_path = 'C:\\Users\\linj\\Desktop\\tmp\\test\\test_less1.docx'
    docx2_path = 'C:\\Users\\linj\\Desktop\\tmp\\test\\test_less2.docx'
    docx1_sentences = extract_sentences_from_docx(docx1_path)
    print(docx1_sentences)
    docx2_sentences = extract_sentences_from_docx(docx2_path)
    print(docx2_sentences)
    cross_matched_pairs = cross_match_sentences(docx1_sentences, docx2_sentences)
    print(cross_matched_pairs)

    # sim = hanlp.load(hanlp.pretrained.sts.STS_ELECTRA_BASE_ZH)
    sim = hanlp.load('C:\\Users\\linj\\AppData\\Roaming\\hanlp\\sts\\sts_electra_base_zh_20210530_200109')

    result = sim(cross_matched_pairs)
    print(result)
    # 找出相似度得分大于0.5的句子对及其分数和具体内容
    high_similarity_pairs_with_scores = [
        [score, pair[0], pair[1]]
        for score, pair in zip(result, cross_matched_pairs)
        if score > 0.5
    ]

    print("相似度得分大于0.5的句子对及其得分：")
    for similarity_score, sentence1, sentence2 in high_similarity_pairs_with_scores:
        print(f"[{similarity_score:.3f}, '{sentence1}', '{sentence2}']")