# -*- coding:utf-8 -*-
# Author: hankcs
# Date: 2021-05-24 13:15
import hanlp

sim = hanlp.load(hanlp.pretrained.sts.STS_ELECTRA_BASE_ZH)
result = sim([
    ['1、对于本项目的理解及重点、难点分析；', '1、对于本工程的理解及重点、难点理解'],
    ['无线路由器怎么无线上网', '2、总体方案及思路'],
    ['北京到上海的动车票', '上海到北京的动车票'],
])
print(type(result))
print(result)
###### pip start########

# pip install python-docx
###### pip end  ########
import re
from docx import Document


def extract_sentences_from_docx(doc_file):
    """
    读取docx文件，返回文件的句子字符串列表
    参数:
    docx_path (str): docx文件的路径
    返回:
    list[str]: 包含文件中所有句子的字符串列表
    """
    # 读取docx文件
    doc = Document(doc_file)
    # 将文档内容合并为一个字符串（按段落连接）
    # text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

    # 按照中文句子进行拆分（这里简单地以句号、问号、感叹号结尾作为句子结束标志）
    # chinese_sentences = re.split(r'(?<=[\u4e00-\u9fa5，。！？;；])(?=\s+)', text)
    chinese_sentences = [match.group(0) for match in re.finditer(r'[^。！!？?;；]*[。！!？?;；]', text)]

    # 过滤掉空字符串
    chinese_sentences = [sentence for sentence in chinese_sentences if sentence]

    # 输出每一句，并统计句子数量
    # sentence_count = 0
    # for index, sentence in enumerate(chinese_sentences):
    #     # print(f"{sentence.strip()}")
    #     print(f"{index + 1}: {sentence.strip()}")
    #     sentence_count += 1
    #
    # print(f"总共识别到 {sentence_count} 句话")
    return chinese_sentences

# 调用函数
# extract_sentences_from_docx('C:\\Users\\linj\\Desktop\\tmp\\test\\test.docx')

